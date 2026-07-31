"""生成合同場景的修訂、批註及雙文件對比 mock。"""

from __future__ import annotations

import argparse
import sys
from copy import deepcopy
from pathlib import Path
from tempfile import TemporaryDirectory
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from lxml import etree

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
XML_NS = "http://www.w3.org/XML/1998/namespace"
W = f"{{{W_NS}}}"


def _w_element(name: str, **attributes: str) -> etree._Element:
    element = etree.Element(f"{W}{name}")
    for key, value in attributes.items():
        element.set(f"{W}{key}", value)
    return element


def _run(text: str, *, deleted: bool = False) -> etree._Element:
    run = _w_element("r")
    text_node = _w_element("delText" if deleted else "t")
    if text[:1].isspace() or text[-1:].isspace():
        text_node.set(f"{{{XML_NS}}}space", "preserve")
    text_node.text = text
    run.append(text_node)
    return run


def _revision(kind: str, revision_id: str, author: str, date: str, text: str) -> etree._Element:
    node = _w_element(kind, id=revision_id, author=author, date=date)
    node.append(_run(text, deleted=kind == "del"))
    return node


def _comment_start(comment_id: str) -> etree._Element:
    return _w_element("commentRangeStart", id=comment_id)


def _comment_end(comment_id: str) -> etree._Element:
    return _w_element("commentRangeEnd", id=comment_id)


def _comment_reference(comment_id: str) -> etree._Element:
    run = _w_element("r")
    properties = _w_element("rPr")
    properties.append(_w_element("rStyle", val="CommentReference"))
    run.append(properties)
    run.append(_w_element("commentReference", id=comment_id))
    return run


def _replace_paragraph_content(paragraph: etree._Element, children: list[etree._Element]) -> None:
    """保留段落屬性，替換其餘 run／修訂／批註標記。"""

    properties = paragraph.find(f"{W}pPr")
    saved_properties = deepcopy(properties) if properties is not None else None
    for child in list(paragraph):
        paragraph.remove(child)
    if saved_properties is not None:
        paragraph.append(saved_properties)
    paragraph.extend(children)


def _comments_xml() -> bytes:
    root = etree.Element(f"{W}comments", nsmap={"w": W_NS, "r": R_NS})
    definitions = [
        ("0", "趙法務", "2025-01-10T09:30:00Z", "請確認合同期限是否需要增加自動續約約定。"),
        ("1", "錢經理", "2025-01-12T15:45:00Z", "請財務確認此付款完成條件是否清晰。"),
    ]
    for comment_id, author, date, content in definitions:
        comment = _w_element("comment", id=comment_id, author=author, date=date, initials=author[:1])
        paragraph = _w_element("p")
        paragraph.append(_run(content))
        comment.append(paragraph)
        root.append(comment)
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def _inject_tracking(base_docx: Path, output_docx: Path) -> None:
    with ZipFile(base_docx) as source:
        files = {name: source.read(name) for name in source.namelist()}

    document_root = etree.fromstring(files["word/document.xml"])
    paragraphs = document_root.xpath(".//w:body/w:p", namespaces={"w": W_NS})
    if len(paragraphs) < 6:
        raise RuntimeError("mock 骨架段落數不足")

    _replace_paragraph_content(paragraphs[0], [_run("合同條款（修訂測試稿）")])
    _replace_paragraph_content(
        paragraphs[1],
        [
            _run("第一條 合同期限為"),
            _comment_start("0"),
            _run("一年"),
            _comment_end("0"),
            _comment_reference("0"),
            _revision("ins", "1", "張三", "2025-01-08T10:00:00Z", "，自2025年1月1日起生效"),
            _run("。"),
        ],
    )
    _replace_paragraph_content(
        paragraphs[2],
        [
            _run("第二條 乙方應在收到合法發票後"),
            _revision("del", "2", "李四", "2025-01-09T11:15:00Z", "七日內"),
            _revision("ins", "3", "陳會計", "2025-01-09T11:20:00Z", "十個工作日內"),
            _comment_start("1"),
            _run("完成付款"),
            _comment_end("1"),
            _comment_reference("1"),
            _run("。"),
        ],
    )
    _replace_paragraph_content(
        paragraphs[3],
        [
            _run("第三條 因本合同引起的爭議，雙方同意提交"),
            _revision("del", "4", "王五", "2025-01-11T14:00:00Z", "北京"),
            _revision("ins", "5", "王五", "2025-01-11T14:01:00Z", "上海"),
            _run("仲裁委員會處理。"),
        ],
    )
    _replace_paragraph_content(paragraphs[4], [_run("第四條 雙方對履約過程中知悉的商業秘密負有保密義務。")])
    _replace_paragraph_content(paragraphs[5], [_run("第五條 違約方應賠償守約方因此遭受的實際損失。")])
    files["word/document.xml"] = etree.tostring(document_root, xml_declaration=True, encoding="UTF-8", standalone=True)
    files["word/comments.xml"] = _comments_xml()

    relationships = etree.fromstring(files["word/_rels/document.xml.rels"])
    existing_ids = {item.get("Id") for item in relationships}
    number = 1
    while f"rId{number}" in existing_ids:
        number += 1
    relationship = etree.SubElement(relationships, f"{{{REL_NS}}}Relationship")
    relationship.set("Id", f"rId{number}")
    relationship.set("Type", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments")
    relationship.set("Target", "comments.xml")
    files["word/_rels/document.xml.rels"] = etree.tostring(relationships, xml_declaration=True, encoding="UTF-8", standalone=True)

    content_types = etree.fromstring(files["[Content_Types].xml"])
    if not content_types.xpath("./ct:Override[@PartName='/word/comments.xml']", namespaces={"ct": CT_NS}):
        override = etree.SubElement(content_types, f"{{{CT_NS}}}Override")
        override.set("PartName", "/word/comments.xml")
        override.set("ContentType", "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml")
    files["[Content_Types].xml"] = etree.tostring(content_types, xml_declaration=True, encoding="UTF-8", standalone=True)

    with ZipFile(output_docx, "w", ZIP_DEFLATED) as target:
        for name, content in files.items():
            target.writestr(name, content)


def _save_clean_document(path: Path, paragraphs: list[str]) -> None:
    document = Document()
    document.add_heading(paragraphs[0], level=1)
    for text in paragraphs[1:]:
        document.add_paragraph(text)
    document.save(path)


def generate_mock(output_dir: Path) -> list[Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    tracked = output_dir / "contract_tracked.docx"
    before = output_dir / "contract_before.docx"
    after = output_dir / "contract_after.docx"

    before_paragraphs = [
        "合同條款（原始版）",
        "第一條 合同期限為一年。",
        "第二條 乙方應在收到合法發票後七日內完成付款。",
        "第三條 因本合同引起的爭議，雙方同意提交北京仲裁委員會處理。",
        "第四條 雙方應每月召開一次履約會議。",
        "第五條 雙方對履約過程中知悉的商業秘密負有保密義務。",
        "第六條 違約方應賠償守約方因此遭受的實際損失。",
    ]
    after_paragraphs = [
        "合同條款（最終版）",
        "第一條 合同期限為一年，自2025年1月1日起生效。",
        "第二條 乙方應在收到合法發票後十個工作日內完成付款。",
        "第三條 因本合同引起的爭議，雙方同意提交上海仲裁委員會處理。",
        "第五條 雙方對履約過程中知悉的商業秘密負有保密義務。",
        "第六條 違約方應賠償守約方因此遭受的實際損失。",
        "第七條 本合同一式兩份，雙方各執一份。",
    ]
    _save_clean_document(before, before_paragraphs)
    _save_clean_document(after, after_paragraphs)

    with TemporaryDirectory() as temporary:
        base = Path(temporary) / "base.docx"
        # 先由 python-docx 建立合法 OPC 骨架，再直接注入底層 WordprocessingML。
        _save_clean_document(base, before_paragraphs[:6])
        _inject_tracking(base, tracked)
    return [tracked, before, after]


def _configure_utf8_console() -> None:
    """避免 Windows CI／舊控制台以 cp1252 輸出中文時崩潰。"""

    for stream in (sys.stdout, sys.stderr):
        reconfigure = getattr(stream, "reconfigure", None)
        if reconfigure is not None:
            try:
                reconfigure(encoding="utf-8", errors="backslashreplace")
            except (AttributeError, OSError):
                pass


def main() -> None:
    _configure_utf8_console()
    parser = argparse.ArgumentParser(description="生成 Word 修訂與對比測試文件")
    parser.add_argument("--output-dir", default="mock_data", help="輸出目錄")
    args = parser.parse_args()
    for path in generate_mock(Path(args.output_dir)):
        print(f"已生成：{path.resolve()}")


if __name__ == "__main__":
    main()
